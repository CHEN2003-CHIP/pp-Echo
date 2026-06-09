from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from pp_agent.attachments.chunkers import (
    chunk_code_file,
    chunk_markdown,
    chunk_plain_text,
    chunk_pdf_pages,
    chunk_structured_json,
    chunk_table_profile,
    source_ref_for_chunk,
)
from pp_agent.attachments.schema import AttachmentChunk, AttachmentKind
from pp_agent.attachments.text_utils import preview_text, read_text_lossy


def extract_attachment(path: Path, *, kind: AttachmentKind, attachment_id: str, session_id: str, filename: str) -> tuple[str, list[AttachmentChunk], dict[str, Any]]:
    """按文件类型抽取文本、生成 chunk 和 inspect metadata，缺少可选依赖时优雅失败。"""

    if kind == AttachmentKind.PDF:
        return _extract_pdf(path, attachment_id=attachment_id, session_id=session_id, filename=filename)
    if kind == AttachmentKind.DOCX:
        return _extract_docx(path, attachment_id=attachment_id, session_id=session_id, filename=filename)
    if kind in {AttachmentKind.TEXT, AttachmentKind.LOG}:
        text = read_text_lossy(path)
        chunks = chunk_plain_text(text, attachment_id=attachment_id, session_id=session_id, filename=filename, kind=kind)
        return text, chunks, {"preview": preview_text(text), "line_count": text.count("\n") + 1}
    if kind == AttachmentKind.MARKDOWN:
        text = read_text_lossy(path)
        chunks = chunk_markdown(text, attachment_id=attachment_id, session_id=session_id, filename=filename)
        headings = [chunk.heading_path for chunk in chunks if chunk.heading_path]
        return text, chunks, {"preview": preview_text(text), "headings": headings[:50]}
    if kind == AttachmentKind.CODE:
        text = read_text_lossy(path)
        chunks, outline = chunk_code_file(text, attachment_id=attachment_id, session_id=session_id, filename=filename)
        return text, chunks, {"preview": preview_text(text), "outline": outline[:200], "line_count": text.count("\n") + 1}
    if kind == AttachmentKind.CSV:
        text = read_text_lossy(path)
        chunks, profile = chunk_table_profile(text, attachment_id=attachment_id, session_id=session_id, filename=filename)
        extracted = json.dumps(profile, ensure_ascii=False, indent=2)
        return extracted, chunks, {"preview": preview_text(extracted), "table": profile}
    if kind in {AttachmentKind.JSON, AttachmentKind.YAML}:
        text = read_text_lossy(path)
        chunks, summary = chunk_structured_json(text, attachment_id=attachment_id, session_id=session_id, filename=filename, kind=kind)
        return text, chunks, {"preview": preview_text(text), "structure": summary}
    return "", [], {"preview": f"{kind.value} attachment is stored but not text indexed.", "indexed": False}


def _extract_pdf(path: Path, *, attachment_id: str, session_id: str, filename: str) -> tuple[str, list[AttachmentChunk], dict[str, Any]]:
    """使用 PyMuPDF 按页抽取 PDF 文本；依赖不存在时返回明确错误。"""

    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PDF parser dependency not installed. Install optional extra 'attachments'.") from exc
    pages: list[tuple[int, str]] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc, start=1):
            pages.append((index, page.get_text("text")))
        page_count = doc.page_count
    text = "\n\n".join(f"# Page {page}\n{body}" for page, body in pages)
    chunks = chunk_pdf_pages(pages, attachment_id=attachment_id, session_id=session_id, filename=filename)
    return text, chunks, {"preview": preview_text(text), "page_count": page_count}


def _extract_docx(path: Path, *, attachment_id: str, session_id: str, filename: str) -> tuple[str, list[AttachmentChunk], dict[str, Any]]:
    """使用 python-docx 抽取段落和表格；依赖不存在时返回明确错误。"""

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("DOCX parser dependency not installed. Install optional extra 'attachments'.") from exc
    document = Document(str(path))
    lines: list[str] = []
    headings: list[list[str]] = []
    current_headings: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.lower().startswith("heading"):
            level = _docx_heading_level(style_name)
            current_headings = current_headings[: level - 1] + [text]
            headings.append(list(current_headings))
            lines.append("#" * min(level, 6) + f" {text}")
        else:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = "\n\n".join(lines)
    chunks = chunk_plain_text(text, attachment_id=attachment_id, session_id=session_id, filename=filename, kind=AttachmentKind.DOCX)
    for chunk in chunks:
        active = _heading_for_range(headings, text, chunk.line_start or 1, chunk.line_end or chunk.line_start or 1)
        if active:
            chunk.heading_path = active
            chunk.section_title = active[-1]
            chunk.source_ref = source_ref_for_chunk(filename, heading_path=active)
    return text, chunks, {"preview": preview_text(text), "paragraph_count": len(document.paragraphs), "table_count": len(document.tables), "headings": headings[:50]}


def _docx_heading_level(style_name: str) -> int:
    """从 python-docx 的样式名中提取 heading 层级，失败时按一级标题处理。"""

    digits = "".join(ch for ch in style_name if ch.isdigit())
    return max(1, min(6, int(digits or "1")))


def _heading_for_line(headings: list[list[str]], text: str, line_no: int) -> list[str]:
    """根据抽取后的 Markdown-like 文本行号找到当前 DOCX heading path。"""

    current: list[str] = []
    heading_index = 0
    for current_line, line in enumerate(text.splitlines(), start=1):
        if heading_index < len(headings) and line.lstrip().startswith("#"):
            current = headings[heading_index]
            heading_index += 1
        if current_line >= line_no:
            return list(current)
    return list(current)


def _heading_for_range(headings: list[list[str]], text: str, line_start: int, line_end: int) -> list[str]:
    """Return the deepest heading active within a DOCX chunk line range."""

    current: list[str] = []
    heading_index = 0
    best: list[str] = []
    for current_line, line in enumerate(text.splitlines(), start=1):
        if heading_index < len(headings) and line.lstrip().startswith("#"):
            current = headings[heading_index]
            heading_index += 1
            if line_start <= current_line <= line_end:
                best = list(current)
        elif line_start <= current_line <= line_end and current:
            best = list(current)
        if current_line >= line_end:
            break
    return best or _heading_for_line(headings, text, line_start)
