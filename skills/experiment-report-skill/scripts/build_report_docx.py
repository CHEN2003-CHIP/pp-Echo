#!/usr/bin/env python3
"""Build a simple DOCX report from a Markdown-style report draft.

This helper is intentionally conservative. It supports headings, paragraphs,
plain Markdown image references, and simple pipe tables. For complex reports,
use a full document generation pipeline or manually polish the generated DOCX.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
except Exception as exc:  # pragma: no cover
    Document = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


def read_text(path: Path) -> str:
    for enc in ["utf-8", "utf-8-sig", "gbk", "latin-1"]:
        try:
            return path.read_text(encoding=enc, errors="replace")
        except Exception:
            continue
    return path.read_text(errors="replace")


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def parse_table(lines: List[str]) -> List[List[str]]:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip Markdown separator row.
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: "Document", rows: List[List[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(max_cols):
            table.cell(i, j).text = row[j] if j < len(row) else ""


def configure_doc(doc: "Document") -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "黑体"


def add_paragraph(doc: "Document", text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25


def build_docx(markdown_path: Path, output_path: Path, base_dir: Path | None = None) -> None:
    if Document is None:
        raise SystemExit(f"python-docx is unavailable: {DOCX_IMPORT_ERROR}\nInstall with: pip install python-docx")

    base_dir = base_dir or markdown_path.parent
    text = read_text(markdown_path)
    lines = text.splitlines()
    doc = Document()
    configure_doc(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if is_table_line(stripped):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        img_match = re.match(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)", stripped)
        if img_match:
            img_path = Path(img_match.group("path"))
            if not img_path.is_absolute():
                img_path = base_dir / img_path
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Cm(13.5))
                alt = img_match.group("alt")
                if alt:
                    cap = doc.add_paragraph(alt)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_paragraph(doc, f"[Image not found: {img_path}]")
            i += 1
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        else:
            add_paragraph(doc, stripped)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a simple DOCX from a Markdown report draft.")
    parser.add_argument("--report_md", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--base_dir", required=False)
    args = parser.parse_args()

    markdown_path = Path(args.report_md).resolve()
    output_path = Path(args.output).resolve()
    base_dir = Path(args.base_dir).resolve() if args.base_dir else markdown_path.parent
    if not markdown_path.exists():
        raise SystemExit(f"Report Markdown not found: {markdown_path}")

    build_docx(markdown_path, output_path, base_dir)
    print(f"Wrote DOCX report: {output_path}")


if __name__ == "__main__":
    main()
